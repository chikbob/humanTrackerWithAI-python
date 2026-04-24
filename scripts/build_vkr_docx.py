from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "vkr"
MARKDOWN_PATH = DOCS_DIR / "VKR_Koltsov_draft.md"
OUTPUT_PATH = ROOT / "ВКР_Кольцов.docx"


def set_page_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def configure_default_font(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.bold = True
    heading1.font.size = Pt(14)
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.bold = True
    heading2.font.size = Pt(14)
    heading2.paragraph_format.first_line_indent = Cm(0)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(6)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_title_page(document: Document) -> None:
    lines = [
        ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, 0),
        ("", False, 0),
        ("на тему", False, 0),
        ("«Разработка веб-системы мониторинга и интеллектуального анализа объектов в реальном времени с использованием нейросетевых моделей»", False, 0),
        ("", False, 0),
        ("", False, 0),
        ("Кольцов", False, 0),
    ]
    for text, bold, before in lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
    document.add_page_break()


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("СОДЕРЖАНИЕ")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    document.add_page_break()


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    text = " ".join(part.strip() for part in buffer if part.strip()).strip()
    buffer.clear()
    if not text:
        return
    p = document.add_paragraph()
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_table_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_label_line(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True


def add_table(document: Document, rows_data: list[list[str]]) -> None:
    header = rows_data[0]
    body = [row for row in rows_data[2:] if len(row) == len(header)]
    table = document.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
    for row_idx, row in enumerate(body, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)
    document.add_paragraph()


def process_markdown(document: Document, markdown_text: str) -> None:
    paragraph_buffer: list[str] = []
    in_code_block = False
    code_buffer: list[str] = []
    lines = markdown_text.splitlines()
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if len(table_buffer) < 2:
            table_buffer = []
            return
        rows = []
        for row in table_buffer:
            parts = [cell.strip() for cell in row.strip().strip("|").split("|")]
            rows.append(parts)
        table_buffer = []
        if len(rows) < 2:
            return
        header = rows[0]
        body = [row for row in rows[2:] if len(row) == len(header)]
        add_table(document, rows)

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_table()
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                add_code_block(document, code_buffer)
                code_buffer = []
            continue
        if in_code_block:
            code_buffer.append(line)
            continue

        if not line.strip():
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            heading_text = heading_match.group(2).strip()
            raw_level = len(heading_match.group(1))
            if raw_level == 1:
                level = 1
            elif raw_level == 2 and re.match(r"^\d+\.\d+", heading_text):
                level = 2
            elif raw_level == 2:
                level = 1
            else:
                level = 2
            add_heading(document, heading_text, level)
            continue

        placeholder_match = re.match(r"^\[МЕСТО ДЛЯ РИСУНКА:\s*(.+?)\]$", line)
        if placeholder_match:
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            continue

        if re.match(r"^Рисунок\s+\d+(\.\d+)?\s+–\s+", line):
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            add_caption(document, line)
            continue

        if line == "Источники, использованные при подготовке обзора:":
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            add_label_line(document, line)
            continue

        if re.match(r"^Таблица\s+\d+(\.\d+)?\.?", line):
            flush_table()
            flush_paragraph(document, paragraph_buffer)
            add_table_caption(document, line)
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_paragraph(document, paragraph_buffer)
            table_buffer.append(line)
            continue
        flush_table()

        if line.startswith("- "):
            flush_paragraph(document, paragraph_buffer)
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.add_run(line[2:].strip())
            continue

        paragraph_buffer.append(line)

    flush_table()
    flush_paragraph(document, paragraph_buffer)


def postprocess_document(document: Document) -> None:
    for p in document.paragraphs:
        text = "".join(r.text for r in p.runs).strip()
        if not text:
            continue
        if re.match(r"^Рисунок\s+\d+(\.\d+)?\s+–\s+", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(12)
        elif re.match(r"^Таблица\s+\d+(\.\d+)?\.?", text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(12)
        elif text == "Источники, использованные при подготовке обзора:":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(14)


def main() -> None:
    markdown_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    marker = "## Реферат"
    if marker in markdown_text:
        markdown_text = markdown_text[markdown_text.index(marker):]
    document = Document()
    set_page_margins(document)
    configure_default_font(document)
    add_title_page(document)
    add_toc(document)
    process_markdown(document, markdown_text)
    postprocess_document(document)
    for section in document.sections:
        add_page_number(section)
    document.save(str(OUTPUT_PATH))
    print(f"saved {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
